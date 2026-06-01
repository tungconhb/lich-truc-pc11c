#!/usr/bin/env python3
"""
Script import dữ liệu từ data.docx lên Firebase Realtime Database.
- Đọc DOCX, trích xuất lịch trực + Trực ban
- Chuyển thành dữ liệu có cấu trúc JSON
- Push lên Firebase (realtime sync cho tất cả người dùng)
- Hỗ trợ chế độ local: tạo data.json dự phòng

Cách dùng:
  python import_docx_final.py          → push Firebase + tạo data.json local
  python import_docx_final.py --local  → chỉ tạo data.json (không push Firebase)

Cấu hình Firebase: sửa biến FIREBASE_CONFIG bên dưới.
"""

import json
import re
import sys
import argparse
from pathlib import Path
from typing import Dict, List, Tuple, Optional

try:
    from docx import Document
except ImportError:
    print("Chưa cài đặt python-docx. Chạy: pip install python-docx")
    sys.exit(1)

try:
    import requests
except ImportError:
    requests = None
    print("⚠️  Chưa cài requests. Firebase push bị tắt.")
    print("   Chạy: pip install requests")


# ==========================================
#        CẤU HÌNH (SỬA TẠI ĐÂY)
# ==========================================

DOCX_PATH = Path(r"C:\Users\Riddler\Desktop\Lịch trực\data.docx")
JSON_PATH = Path(r"C:\Users\Riddler\Desktop\Lịch trực\data.json")
HTML_PATH = Path(r"C:\Users\Riddler\Desktop\Lịch trực\pc11c.html")

# Firebase Realtime Database URL (lấy từ Firebase Console)
FIREBASE_DB_URL = "https://lich-truc-pc11c-default-rtdb.asia-southeast1.firebasedatabase.app"

# Firebase Database Secret (lấy từ Console → Project Settings → Service Accounts → Database Secrets)
FIREBASE_SECRET = "Ax61a96jv7As9WIMOgGLDHBPdFi4TzXDZbZzVI1H"

# Thứ tự section mặc định (không phụ thuộc vào thứ tự trong file DOCX)
DEFAULT_SECTION_ORDER = [
    "LÃNH ĐẠO",
    "ĐỘI THAM MƯU - HẬU CẦN",
    "ĐỘI QUẢN GIÁO",
    "ĐỘI CẢNH SÁT BẢO VỆ",
    "PHÂN TRẠI QUẢN LÝ PHẠM NHÂN",
    "PHÂN TRẠI QUẢN LÝ PHẠM NHÂN (Tổ Cao Phong)",
    "PHÂN TRẠI LƯƠNG SƠN",
    "PHÂN TRẠI MAI CHÂU",
    "PHÂN TRẠI LẠC THUỶ",
]

# ==========================================


def parse_header_dates(doc: Document) -> List[str]:
    """Đọc ngày tháng từ header Table 1 DOCX."""
    if len(doc.tables) < 2:
        return []

    table = doc.tables[1]
    header_row = table.rows[0]
    cells = [cell.text.strip() for cell in header_row.cells]

    dates = []
    for cell in cells[2:9]:
        m = re.search(r'\(Ngày\s*(\d{1,2}/\d{1,2}/\d{4})\)', cell)
        if m:
            dates.append(m.group(1))
        else:
            m = re.search(r'(\d{1,2}/\d{1,2}/\d{4})', cell)
            dates.append(m.group(1) if m else "")

    return dates


def extract_truc_ban(doc: Document) -> Dict[str, str]:
    """Trích xuất Trực ban từ phần 'Trực ban đơn vị' trong DOCX."""
    truc_ban = {}
    full_text = "\n".join([p.text for p in doc.paragraphs])

    patterns = {
        "25/5": r"Ngày 25/5/\d{4}:\s*(.+)",
        "26/5": r"Ngày 26/5/\d{4}:\s*(.+)",
        "27/5": r"Ngày 27/5/\d{4}:\s*(.+)",
        "28/5": r"Ngày 28/5/\d{4}:\s*(.+)",
        "29/5": r"Ngày 29/5/\d{4}:\s*(.+)",
        "30/5": r"Ngày 30/5/\d{4}:\s*(.+)",
        "31/5": r"Ngày 31/5/\d{4}:\s*(.+)",
    }
    for day, pat in patterns.items():
        m = re.search(pat, full_text)
        if m:
            truc_ban[day] = m.group(1).strip()

    return truc_ban


def parse_docx_table(doc: Document) -> List[Dict]:
    """
    Phân tích bảng chính (Table 1) trong DOCX.
    Trả về danh sách sections với cấu trúc phù hợp cho Firebase.
    """
    if len(doc.tables) < 2:
        raise ValueError("DOCX không có bảng chính (cần ít nhất 2 bảng)")

    table = doc.tables[1]
    sections = []
    current_section = None
    current_rows = []

    for row_idx, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]

        if not any(cells):
            continue

        if cells[0] == "STT":
            continue

        # Hàng 1: LÃNH ĐẠO
        if row_idx == 1 and len(cells) >= 9:
            days = cells[2:9]
            days = days + [''] * (7 - len(days))
            sections.append({
                "name": "LÃNH ĐẠO",
                "rows": [{"nhiem_vu": cells[1] if len(cells) > 1 else "Trực lãnh đạo", "days": days}]
            })
            continue

        # Section header: tất cả ô giống nhau
        unique = set(c for c in cells if c)
        if len(unique) == 1 and len(cells) >= 2:
            section_name = list(unique)[0]
            if current_section:
                sections.append({"name": current_section, "rows": current_rows})
            current_section = section_name
            current_rows = []
            continue

        # Data row
        if current_section and len(cells) >= 9:
            days = cells[2:9]
            days = days + [''] * (7 - len(days))
            current_rows.append({"nhiem_vu": cells[1] if len(cells) > 1 else "", "days": days})

    # Lưu section cuối
    if current_section:
        sections.append({"name": current_section, "rows": current_rows})

    return sections


def get_html_template_rows(html_path: Path) -> Dict[str, int]:
    """Đọc HTML template để lấy số dòng mỗi section."""
    if not html_path.exists():
        return {}

    with open(html_path, 'r', encoding='utf-8') as f:
        html = f.read()

    tbody_match = re.search(r'<tbody[^>]*>(.*?)</tbody>', html, re.DOTALL)
    if not tbody_match:
        return {}

    tbody = tbody_match.group(1)
    template = {}
    current_section = None
    row_count = 0

    for line in tbody.split('\n'):
        line = line.strip()
        sm = re.search(r'<tr class="section-row"><td colspan="10">(.*?)</td></tr>', line)
        if sm:
            if current_section is not None:
                template[current_section] = row_count
            current_section = sm.group(1)
            row_count = 0
        elif line.startswith('<tr'):
            row_count += 1

    if current_section is not None:
        template[current_section] = row_count

    return template


def match_docx_to_template(
    sections: List[Dict],
    template: Dict[str, int]
) -> List[Dict]:
    """
    Căn chỉnh số dòng DOCX khớp với template HTML.
    - DOCX ít hơn template: thêm dòng trống
    - DOCX nhiều hơn template: giữ tất cả dòng DOCX
    - Luôn sắp xếp theo DEFAULT_SECTION_ORDER
    """
    result = []
    docx_map = {normalize_name(s["name"]): s for s in sections}

    # Sắp xếp template theo DEFAULT_SECTION_ORDER
    ordered_template = {}
    for section_name in DEFAULT_SECTION_ORDER:
        norm_name = normalize_name(section_name)
        for t_name, t_rows in template.items():
            if normalize_name(t_name) == norm_name:
                ordered_template[section_name] = t_rows
                break
    
    # Thêm các section trong template nhưng không có trong order
    for section_name, target_rows in template.items():
        if section_name not in ordered_template:
            ordered_template[section_name] = target_rows

    unmatched_sections = [s for s in sections if normalize_name(s["name"]) not in template]

    for section_name, target_rows in ordered_template.items():
        norm_name = normalize_name(section_name)
        if norm_name in docx_map:
            sec = docx_map[norm_name]
            rows = sec["rows"]
            if len(rows) < target_rows:
                rows = rows + [{"nhiem_vu": "", "days": [''] * 7}] * (target_rows - len(rows))
            result.append({"name": section_name, "rows": rows})
        else:
            # Không có dữ liệu DOCX → tạo section trống
            result.append({"name": section_name, "rows": [{"nhiem_vu": "", "days": [''] * 7}] * target_rows})

    # Thêm các section có trong DOCX nhưng không có trong template
    for sec in unmatched_sections:
        if sec["name"] not in [r["name"] for r in result]:
            result.append(sec)

    return result


def normalize_name(name: str) -> str:
    """Chuẩn hóa tên section để khớp giữa DOCX và HTML."""
    name = name.upper()
    name = name.replace("THUỶ", "THỦY")
    name = name.replace("(TỔ CAO PHONG)", "(Tổ Cao Phong)")
    return name.strip()


def build_duty_info(truc_ban: Dict[str, str]) -> str:
    """Tạo chuỗi HTML Trực ban."""
    day_order = [
        "25/5/2026", "26/5/2026", "27/5/2026",
        "28/5/2026", "29/5/2026", "30/5/2026", "31/5/2026"
    ]
    items = []
    for day in day_order:
        day_key = day.split('/')[0] + '/' + day.split('/')[1]
        if day_key in truc_ban:
            items.append(f'<strong>{day}</strong>: {truc_ban[day_key]}')

    return '★ Trực ban:&ensp;' + '&ensp;'.join(items) + '.'


def push_to_firebase(data: Dict) -> bool:
    """Đẩy dữ liệu lên Firebase Realtime Database."""
    if not requests:
        print("❌ Thiếu thư viện requests. Chạy: pip install requests")
        return False

    if "THAY_BANG" in FIREBASE_DB_URL:
        print("❌ Chưa cấu hình FIREBASE_DB_URL. Sửa trong file import_docx_final.py")
        return False

    if "THAY_BANG" in FIREBASE_SECRET:
        print("❌ Chưa cấu hình FIREBASE_SECRET. Sửa trong file import_docx_final.py")
        return False

    # Lấy dữ liệu hiện tại từ Firebase
    try:
        get_url = FIREBASE_DB_URL.rstrip('/') + '/.json?auth=' + FIREBASE_SECRET
        resp = requests.get(get_url, timeout=10)
        if resp.status_code != 200:
            print(f"❌ Không lấy được dữ liệu hiện tại: {resp.status_code}")
            return False
        current_data = resp.json() or {}
    except Exception as e:
        print(f"❌ Lỗi kết nối Firebase: {e}")
        return False

    # Merge dữ liệu: chỉ cập nhật những section có trong DOCX
    # Giữ nguyên các section khác
    if "schedule" in current_data and "sections" in current_data["schedule"]:
        current_sections = current_data["schedule"]["sections"]
        new_sections = data.get("schedule", {}).get("sections", [])
        
        # Tạo map từ DOCX
        docx_map = {normalize_name(s["name"]): s for s in new_sections}
        
        # Merge: cập nhật section có trong DOCX, giữ nguyên section không có
        merged_sections = []
        for sec in current_sections:
            norm_name = normalize_name(sec["name"])
            if norm_name in docx_map:
                # Cập nhật từ DOCX
                merged_sections.append(docx_map[norm_name])
                del docx_map[norm_name]
            else:
                # Giữ nguyên dữ liệu cũ
                merged_sections.append(sec)
        
        # Thêm các section mới từ DOCX (nếu có)
        for sec in new_sections:
            norm_name = normalize_name(sec["name"])
            if norm_name in docx_map:
                merged_sections.append(sec)
        
        # Sắp xếp lại theo DEFAULT_SECTION_ORDER
        final_sections = []
        for name in DEFAULT_SECTION_ORDER:
            for sec in merged_sections:
                if normalize_name(sec["name"]) == normalize_name(name):
                    final_sections.append(sec)
                    break
        # Thêm các section còn lại
        for sec in merged_sections:
            if sec not in final_sections:
                final_sections.append(sec)
        
        data["schedule"]["sections"] = final_sections
    
    # Cập nhật headerDates và dutyInfo nếu có
    if "headerDates" in data:
        current_data["headerDates"] = data["headerDates"]
    if "dutyInfo" in data:
        current_data["dutyInfo"] = data["dutyInfo"]
    
    # Gửi dữ liệu đã merge
    put_url = FIREBASE_DB_URL.rstrip('/') + '/.json?auth=' + FIREBASE_SECRET
    try:
        resp = requests.put(put_url, json=current_data, timeout=30)
        if resp.status_code == 200:
            print("✅ Đã merge dữ liệu lên Firebase (giữ nguyên dữ liệu cũ)")
            return True
        else:
            print(f"❌ Firebase trả về lỗi {resp.status_code}: {resp.text[:200]}")
            return False
    except Exception as e:
        print(f"❌ Lỗi gửi dữ liệu: {e}")
        return False


def generate_local_json(data: Dict):
    """Tạo file data.json dự phòng (tương thích với phiên bản cũ)."""
    sections = data.get("schedule", {}).get("sections", [])
    rows = []
    stt = 1

    for sec in sections:
        rows.append(f'<tr class="section-row"><td colspan="10">{sec["name"]}</td></tr>')
        for r in sec.get("rows", []):
            days_html = ''.join(f'<td class="col-day">{d}</td>' for d in (r.get("days", []) + [''] * 7)[:7])
            rows.append(f'<tr><td>{stt}</td><td class="col-nv">{r.get("nhiem_vu", "")}</td>{days_html}</tr>')
            stt += 1

    with open(JSON_PATH, 'w', encoding='utf-8') as f:
        json.dump({"html": '\n'.join(rows)}, f, ensure_ascii=False)

    print(f"✅ Đã tạo data.json dự phòng ({JSON_PATH})")


def update_html_header(dates: List[str]) -> bool:
    """Cập nhật header ngày tháng trong pc11c.html."""
    if not HTML_PATH.exists() or not dates or len(dates) < 7:
        return False

    with open(HTML_PATH, 'r', encoding='utf-8') as f:
        html = f.read()

    day_names = ["Thứ 2", "Thứ 3", "Thứ 4", "Thứ 5", "Thứ 6", "Thứ 7", "Chủ Nhật"]
    for day_name, date_str in zip(day_names, dates):
        html = re.sub(
            rf'>{re.escape(day_name)}<br>\([^)]*\)</th>',
            f'>{day_name}</th>',
            html
        )
        old = f">{day_name}</th>"
        new = f">{day_name}<br>({date_str})</th>"
        html = html.replace(old, new)

    with open(HTML_PATH, 'w', encoding='utf-8') as f:
        f.write(html)

    return True


def main():
    parser = argparse.ArgumentParser(description='Import DOCX → Firebase / Local')
    parser.add_argument('--local', action='store_true', help='Chỉ tạo data.json, không push Firebase')
    parser.add_argument('--firebase', action='store_true', help='Chỉ push Firebase, không tạo data.json')
    args = parser.parse_args()

    print("=" * 60)
    print("  IMPORT DOCX → FIREBASE REALTIME DATABASE")
    print("=" * 60)

    # 1. Đọc DOCX
    if not DOCX_PATH.exists():
        print(f"❌ Không tìm thấy {DOCX_PATH}")
        sys.exit(1)

    doc = Document(str(DOCX_PATH))
    sections_raw = parse_docx_table(doc)
    truc_ban = extract_truc_ban(doc)
    header_dates = parse_header_dates(doc)

    total_rows = sum(len(s["rows"]) for s in sections_raw)
    print(f"\n📊 DOCX: {len(sections_raw)} sections, {total_rows} dòng")
    for s in sections_raw:
        print(f"   {s['name']}: {len(s['rows'])} dòng")

    # 2. Căn chỉnh với HTML template
    template = get_html_template_rows(HTML_PATH)
    if template:
        sections = match_docx_to_template(sections_raw, template)
        print(f"\n📋 Template HTML: {len(template)} sections")
        print(f"📈 Sau căn chỉnh: {len(sections)} sections, {sum(len(s['rows']) for s in sections)} dòng")
        # Đảm bảo thứ tự cuối cùng theo DEFAULT_SECTION_ORDER
        final_sections = []
        for name in DEFAULT_SECTION_ORDER:
            for sec in sections:
                if normalize_name(sec["name"]) == normalize_name(name):
                    final_sections.append(sec)
                    break
        # Thêm các section còn lại
        for sec in sections:
            if sec not in final_sections:
                final_sections.append(sec)
        sections = final_sections
    else:
        print("\n⚠️  Không đọc được template HTML, dùng dữ liệu DOCX gốc")
        # Sắp xếp sections_raw theo DEFAULT_SECTION_ORDER
        sections = []
        for name in DEFAULT_SECTION_ORDER:
            for sec in sections_raw:
                if normalize_name(sec["name"]) == normalize_name(name):
                    sections.append(sec)
                    break
        # Thêm các section còn lại
        for sec in sections_raw:
            if sec not in sections:
                sections.append(sec)

    # 3. Xây dựng dữ liệu Firebase
    duty_info = build_duty_info(truc_ban)
    firebase_data = {
        "schedule": {"sections": sections},
        "dutyInfo": duty_info,
        "headerDates": header_dates[:7] if len(header_dates) >= 7 else []
    }

    # 4. Push Firebase
    firebase_ok = False
    if not args.local:
        print("\n☁️  Đang push lên Firebase...")
        firebase_ok = push_to_firebase(firebase_data)

    # 5. Tạo data.json dự phòng
    if not args.firebase:
        generate_local_json(firebase_data)

    # 6. Cập nhật header HTML
    if header_dates and len(header_dates) >= 7:
        if update_html_header(header_dates):
            print("✅ Đã cập nhật header HTML với ngày tháng")

    # 7. Hiển thị Trực ban
    if truc_ban:
        print("\n📅 Trực ban đơn vị:")
        for day in ["25/5", "26/5", "27/5", "28/5", "29/5", "30/5", "31/5"]:
            if day in truc_ban:
                print(f"   {day}: {truc_ban[day]}")

    print("\n" + "=" * 60)
    if firebase_ok:
        print("🎉 DỮ LIỆU ĐÃ CÓ TRÊN FIREBASE - MỌI NGƯỜI CÓ THỂ XEM NGAY")
    elif args.local:
        print("📁 Đã tạo data.json local. Dùng --firebase để push lên cloud.")
    else:
        print("⚠️  Firebase chưa được cấu hình. Sửa FIREBASE_DB_URL và FIREBASE_SECRET.")
    print("=" * 60)


if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print(f"\n❌ Lỗi: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)